import io
from datetime import datetime

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


def build_markdown(meeting: dict) -> bytes:
    tasks_lines = "\n".join(
        f"- **{t.get('task', '')}** — Responsável: {t.get('owner', '—')} | Prazo: {t.get('deadline', '—')}"
        for t in (meeting.get("tarefas") or [])
    ) or "Nenhuma tarefa identificada."

    decisions_lines = "\n".join(
        f"- {d}" for d in (meeting.get("decisoes") or [])
    ) or "Nenhuma decisão identificada."

    risks_lines = "\n".join(
        f"- {r}" for r in (meeting.get("riscos") or [])
    ) or "Nenhum risco identificado."

    content = f"""# {meeting['title']}

## Resumo Executivo
{meeting['resumo_executivo']}

## Ata Formal
{meeting['ata_formal']}

## Tarefas
{tasks_lines}

## Decisões
{decisions_lines}

## Riscos
{risks_lines}

---
Gerado em {datetime.now().isoformat()}
"""
    return content.encode("utf-8")


def build_docx(meeting: dict) -> bytes:
    doc = Document()
    doc.add_heading(meeting["title"], 0)
    doc.add_heading("Resumo Executivo", level=1)
    doc.add_paragraph(meeting["resumo_executivo"])
    doc.add_heading("Ata Formal", level=1)
    doc.add_paragraph(meeting["ata_formal"])
    doc.add_heading("Tarefas", level=1)
    for task in meeting["tarefas"]:
        doc.add_paragraph(f"- {task.get('task')} | Responsável: {task.get('owner')} | Prazo: {task.get('deadline')}")

    doc.add_heading("Decisões", level=1)
    for item in meeting["decisoes"]:
        doc.add_paragraph(f"- {item}")

    doc.add_heading("Riscos", level=1)
    for item in meeting["riscos"]:
        doc.add_paragraph(f"- {item}")

    stream = io.BytesIO()
    doc.save(stream)
    return stream.getvalue()


def build_pdf(meeting: dict) -> bytes:
    stream = io.BytesIO()
    pdf = canvas.Canvas(stream, pagesize=A4)
    width, height = A4
    y = height - 50

    def write_line(text: str, top: int = 15):
        nonlocal y
        if y < 50:
            pdf.showPage()
            y = height - 50
        pdf.drawString(40, y, str(text)[:120])
        y -= top

    def write_section(title: str, lines: list[str]):
        write_line("")
        write_line(title, 18)
        for line in lines:
            write_line(f"  {line}")

    write_line(meeting["title"], 22)

    write_section("Resumo Executivo:", meeting["resumo_executivo"].split("\n"))
    write_section("Ata Formal:", meeting["ata_formal"].split("\n"))

    write_section(
        "Tarefas:",
        [
            f"- {t.get('task', '')} | Resp: {t.get('owner', '—')} | Prazo: {t.get('deadline', '—')}"
            for t in (meeting.get("tarefas") or [])
        ] or ["Nenhuma tarefa identificada."],
    )

    write_section(
        "Decisões:",
        [f"- {d}" for d in (meeting.get("decisoes") or [])] or ["Nenhuma decisão identificada."],
    )

    write_section(
        "Riscos:",
        [f"- {r}" for r in (meeting.get("riscos") or [])] or ["Nenhum risco identificado."],
    )

    pdf.save()
    stream.seek(0)
    return stream.read()
